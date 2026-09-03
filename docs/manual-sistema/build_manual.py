from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUTPUT = ROOT / "Manual_do_Usuario_Metalique_Infinity.docx"

RED = RGBColor(218, 15, 15)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(89, 89, 89)
LIGHT_RED = "FDECEC"


SCREENS = [
    {
        "section": "2 ACESSO AO SISTEMA",
        "title": "2.1 Tela inicial",
        "image": "01-inicio.png",
        "purpose": "Apresentar o Metalique Infinity e direcionar o visitante para a entrada no sistema ou para a solicitação de acesso.",
        "access": "Abra o endereço principal do sistema no navegador.",
        "steps": [
            "Selecione Log-in para entrar com uma conta existente.",
            "Selecione Solicitar acesso quando ainda não possuir uma conta.",
            "Use Home, Ajuda e Contato para navegar pelas informações públicas.",
        ],
        "note": "Os botões disponíveis podem mudar quando o usuário já estiver autenticado.",
    },
    {
        "title": "2.2 Tela de login",
        "image": "02-login.png",
        "purpose": "Autenticar o usuário e liberar somente as áreas permitidas para sua conta.",
        "access": "Na tela inicial, selecione Log-in.",
        "steps": [
            "Informe o e-mail cadastrado.",
            "Digite a senha; use o ícone de olho para mostrar ou ocultar os caracteres.",
            "Selecione Entrar e aguarde o redirecionamento.",
        ],
        "note": "Em caso de mensagem de credenciais inválidas, confira o e-mail, a senha e se a conta está ativa.",
    },
    {
        "title": "2.3 Recuperação de senha",
        "image": "22-recuperar-senha.png",
        "purpose": "Solicitar ao administrador autorização para cadastrar uma nova senha.",
        "access": "Na tela de login, selecione Esqueceu sua senha?.",
        "steps": [
            "Informe o e-mail da conta.",
            "Selecione Solicitar recuperação.",
            "Aguarde a análise do administrador; após a aprovação, retorne à tela para cadastrar a nova senha.",
        ],
        "note": "A nova senha deve ter de 8 a 72 caracteres, pelo menos um número e um caractere especial.",
    },
    {
        "title": "2.4 Solicitação de acesso",
        "image": "03-solicitar-acesso.png",
        "purpose": "Registrar o pedido de criação de acesso para um colaborador.",
        "access": "Na tela inicial, selecione Solicitar acesso.",
        "steps": [
            "Preencha nome completo, setor e cargo.",
            "Informe a data de admissão.",
            "Selecione Solicitar acesso e aguarde a confirmação.",
        ],
        "note": "O pedido permanece pendente até ser analisado por um administrador.",
    },
    {
        "section": "3 ÁREA AUTENTICADA",
        "title": "3.1 Dashboard",
        "image": "04-dashboard.png",
        "purpose": "Funcionar como página inicial da área interna e ponto de partida para os módulos liberados.",
        "access": "Após o login, selecione Dashboard no menu superior.",
        "steps": [
            "Use o menu superior para alternar entre Dashboard, Qualidade e Usuários.",
            "Use a lupa para pesquisar páginas e ações.",
            "Use o sino para acompanhar notificações e o avatar para abrir a conta.",
        ],
        "note": "O menu mostra apenas módulos permitidos para o perfil autenticado.",
    },
    {
        "title": "3.2 Pesquisa global",
        "image": "05-pesquisa-global.png",
        "purpose": "Localizar rapidamente módulos, páginas e ações do sistema.",
        "access": "Selecione o botão de lupa no cabeçalho.",
        "steps": [
            "Digite uma palavra relacionada ao que deseja abrir.",
            "Confira os resultados exibidos.",
            "Selecione o resultado desejado para navegar diretamente até ele.",
        ],
        "note": "A pesquisa respeita as permissões da conta e não exibe áreas bloqueadas.",
    },
    {
        "title": "3.3 Notificações",
        "image": "06-notificacoes.png",
        "purpose": "Centralizar solicitações e atualizações que exigem conhecimento ou ação do usuário.",
        "access": "Selecione o sino no cabeçalho.",
        "steps": [
            "Leia o título, a descrição e a data de cada item.",
            "Selecione uma notificação para abrir a ação relacionada, quando disponível.",
            "Use Marcar como lidas para remover o destaque dos itens já verificados.",
        ],
        "note": "Administradores podem receber solicitações de acesso e de recuperação de senha.",
    },
    {
        "title": "3.4 Perfil - dados pessoais",
        "image": "07-perfil.png",
        "purpose": "Permitir a consulta e a atualização dos dados pessoais e da foto do usuário.",
        "access": "Selecione o avatar no cabeçalho e abra Meu perfil.",
        "steps": [
            "Abra a aba Dados pessoais.",
            "Atualize os campos permitidos e, se necessário, escolha uma nova foto.",
            "Salve as alterações e confira a mensagem de confirmação.",
        ],
        "note": "A foto aceita arquivos JPG, PNG ou WebP conforme as validações mostradas na tela.",
    },
    {
        "title": "3.5 Perfil - alteração de senha",
        "image": "08-alterar-senha.png",
        "purpose": "Trocar a senha de uma conta autenticada.",
        "access": "Em Meu perfil, selecione a aba Senha.",
        "steps": [
            "Informe a senha atual.",
            "Digite a nova senha e repita-a no campo de confirmação.",
            "Selecione Alterar senha.",
        ],
        "note": "A nova senha deve ser diferente da atual e cumprir a política de segurança do sistema.",
    },
    {
        "title": "3.6 Menu da conta",
        "image": "21-menu-conta.png",
        "purpose": "Reunir os atalhos pessoais e o encerramento seguro da sessão.",
        "access": "Selecione a seta ao lado do avatar no cabeçalho.",
        "steps": [
            "Abra Meu perfil para editar dados pessoais ou senha.",
            "Confira o nome, o cargo e o tipo de conta em uso.",
            "Selecione Sair ao terminar o trabalho, principalmente em computadores compartilhados.",
        ],
        "note": "Fechar apenas a aba do navegador não substitui o comando Sair.",
    },
    {
        "section": "4 MÓDULO QUALIDADE",
        "title": "4.1 Visão geral e aba RAPs",
        "image": "09-qualidade-raps.png",
        "purpose": "Apresentar os indicadores gerais dos relatórios de ação preventiva e corretiva.",
        "access": "Selecione Qualidade no menu superior e, depois, RAPs.",
        "steps": [
            "Consulte os cartões de total, período, clientes/lotes e modelo mais recorrente.",
            "Analise os gráficos por mês, tipo de problema e código atribuído.",
            "Selecione barras, pontos ou fatias para destacar o subconjunto nos demais gráficos.",
        ],
        "note": "O botão Tabela mostra os mesmos dados do gráfico em formato tabular.",
    },
    {
        "title": "4.2 Filtros da Qualidade",
        "image": "16-filtros-qualidade.png",
        "purpose": "Aplicar o mesmo recorte a todos os indicadores e registros do módulo.",
        "access": "Na tela Qualidade, selecione Filtros.",
        "steps": [
            "Escolha ano, mês, barracão, gate ou outro critério necessário.",
            "Combine filtros para restringir a análise.",
            "Use Limpar para retornar à visão completa e Fechar para recolher o painel.",
        ],
        "note": "Filtros ativos afetam as abas, os totais, as tabelas e os registros.",
    },
    {
        "title": "4.3 Aba Unidades",
        "image": "10-qualidade-unidades.png",
        "purpose": "Comparar ocorrências por barracão, gate e unidade operacional.",
        "access": "No módulo Qualidade, selecione Unidades.",
        "steps": [
            "Confira os cartões de resumo da distribuição por unidade.",
            "Use os gráficos para identificar barracões e gates com maior recorrência.",
            "Selecione um elemento do gráfico para aprofundar o recorte.",
        ],
        "note": "Interprete os valores junto com o volume de trabalho de cada unidade.",
    },
    {
        "title": "4.4 Aba Produtos",
        "image": "11-qualidade-produtos.png",
        "purpose": "Analisar os apontamentos por tipo de máquina e modelo.",
        "access": "No módulo Qualidade, selecione Produtos.",
        "steps": [
            "Identifique os tipos de máquina com maior número de apontamentos.",
            "Compare os modelos e os códigos predominantes.",
            "Use a seleção nos gráficos para investigar um produto específico.",
        ],
        "note": "Os indicadores refletem o filtro global aplicado no módulo.",
    },
    {
        "title": "4.5 Aba Produtos Coletados",
        "image": "12-qualidade-produtos-coletados.png",
        "purpose": "Acompanhar os registros de coleta e expedição de máquinas.",
        "access": "No módulo Qualidade, selecione Produtos Coletados.",
        "steps": [
            "Consulte o total de coletas e sua evolução por período.",
            "Compare os registros por cliente, máquina ou modelo.",
            "Use os gráficos e tabelas para verificar concentração ou tendência.",
        ],
        "note": "O número do registro de coleta utiliza o prefixo RETIR.",
    },
    {
        "title": "4.6 Aba Colaboradores",
        "image": "13-qualidade-colaboradores.png",
        "purpose": "Analisar a participação dos colaboradores nos RAPs e apoiar ações de orientação.",
        "access": "No módulo Qualidade, selecione Colaboradores.",
        "steps": [
            "Consulte o ranking de participações.",
            "Selecione um colaborador para destacar seus códigos e sua evolução.",
            "Use a informação para direcionar instruções específicas, sem tratá-la isoladamente como ranking de desempenho.",
        ],
        "note": "Um RAP pode envolver até três colaboradores; por isso, participações e quantidade de RAPs não são medidas idênticas.",
    },
    {
        "title": "4.7 Aba Qualidade",
        "image": "14-qualidade-satisfacao.png",
        "purpose": "Acompanhar satisfação, reclamações e relação entre coletas e ocorrências de clientes.",
        "access": "No módulo Qualidade, selecione Qualidade.",
        "steps": [
            "Confira as taxas de satisfação e reclamação.",
            "Compare coletas e reclamações por mês.",
            "Consulte a lista de reclamações para avaliar a ocorrência e definir tratativas.",
        ],
        "note": "A taxa de satisfação é calculada a partir das coletas sem reclamação no período analisado.",
    },
    {
        "title": "4.8 Aba Registros",
        "image": "15-qualidade-registros.png",
        "purpose": "Consultar, paginar, imprimir e, quando permitido, excluir RAPs e coletas existentes.",
        "access": "No módulo Qualidade, selecione Registros.",
        "steps": [
            "Localize o RAP ou RETIR pela tabela e pelos filtros.",
            "Use Imprimir para abrir a folha do registro e a caixa de impressão do navegador.",
            "Use Anterior e Próxima para navegar; a opção Excluir aparece apenas para contas autorizadas.",
        ],
        "note": "A exclusão é permanente e exige confirmação. Registros relacionados devem ser conferidos antes da ação.",
    },
    {
        "title": "4.9 Novo apontamento (RAP)",
        "image": "17-novo-rap.png",
        "purpose": "Registrar uma não conformidade, sua origem, os envolvidos e a ação imediata.",
        "access": "Na tela Qualidade, selecione Novo RAP.",
        "steps": [
            "Informe data, identificação, cliente/lote, máquina, área, gate, local, código e descrição.",
            "Selecione de um a três colaboradores envolvidos.",
            "Registre a ação imediata e indique se o checklist precisa ser atualizado.",
            "Selecione Gravar apontamento; o número RAP é gerado automaticamente.",
        ],
        "note": "Campos com asterisco são obrigatórios e a descrição deve ter no mínimo 10 caracteres.",
    },
    {
        "title": "4.10 Novo produto coletado",
        "image": "18-nova-coleta.png",
        "purpose": "Registrar a coleta/expedição de uma máquina, responsáveis, ocorrências e evidências fotográficas.",
        "access": "Na tela Qualidade, selecione Nova coleta.",
        "steps": [
            "Informe data, cliente, tipo de máquina, modelo e ocorrências do carregamento.",
            "Selecione o responsável e até dois colaboradores adicionais.",
            "Adicione de uma a seis fotos do carregamento, com até 5 MB por arquivo.",
            "Registre a ação imediata, informe se o formulário precisa ser alterado e selecione Gravar coleta.",
        ],
        "note": "São aceitas imagens JPEG, PNG ou WebP. O número RETIR é criado na gravação.",
    },
    {
        "section": "5 ADMINISTRAÇÃO DE USUÁRIOS",
        "title": "5.1 Lista de usuários",
        "image": "19-usuarios.png",
        "purpose": "Consultar contas, cargos, setores, tipos, permissões e status de acesso.",
        "access": "Selecione Usuários no menu superior. É necessária a permissão Administrar usuários.",
        "steps": [
            "Confira os dados e o status de cada conta.",
            "Use o lápis para editar uma conta permitida.",
            "Use a lixeira para iniciar a exclusão, quando o botão estiver habilitado.",
        ],
        "note": "A conta administradora principal é protegida e não pode ser desativada ou excluída pela interface.",
    },
    {
        "title": "5.2 Novo usuário e permissões",
        "image": "20-novo-usuario.png",
        "purpose": "Criar contas e definir exatamente os módulos e ações disponíveis.",
        "access": "Na tela Usuários, selecione Novo usuário.",
        "steps": [
            "Preencha nome, cargo, setor, e-mail, tipo de conta e senha inicial.",
            "Para uma conta comum, marque as permissões necessárias em Geral, Qualidade e Administração.",
            "Mantenha Conta ativa marcada e selecione Salvar usuário.",
        ],
        "note": "Administradores possuem acesso total. Contas comuns exibem somente as opções concedidas e podem exigir troca de senha no primeiro acesso.",
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 0) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=12, bold=False, color=BLACK, italic=False) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=10)


def configure_page(section, numbered=False) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    if numbered:
        section.header.is_linked_to_previous = False
        header_p = section.header.paragraphs[0]
        header_p.clear()
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(header_p)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1.25)

    for name, size, before, after, color in (
        ("Heading 1", 14, 18, 10, BLACK),
        ("Heading 2", 13, 14, 8, BLACK),
        ("Heading 3", 12, 10, 6, GRAY),
    ):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Cm(0)

    for name in ("List Number", "List Bullet"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Cm(0.75)
        style.paragraph_format.first_line_indent = Cm(-0.5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.15

    if "Figure Caption" not in styles:
        cap = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Figure Caption"]
    cap.font.name = "Times New Roman"
    cap._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    cap.font.size = Pt(10)
    cap.font.color.rgb = BLACK
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(2)
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.keep_with_next = True
    cap.paragraph_format.first_line_indent = Cm(0)

    if "Figure Source" not in styles:
        src = styles.add_style("Figure Source", WD_STYLE_TYPE.PARAGRAPH)
    else:
        src = styles["Figure Source"]
    src.font.name = "Times New Roman"
    src._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    src._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    src.font.size = Pt(10)
    src.font.color.rgb = GRAY
    src.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.paragraph_format.space_after = Pt(8)
    src.paragraph_format.line_spacing = 1.0
    src.paragraph_format.first_line_indent = Cm(0)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("METALIQUE")
    set_run_font(r, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(115)
    r = p.add_run("MANUAL DO USUÁRIO")
    set_run_font(r, size=22, bold=True, color=RED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Sistema Metalique Infinity")
    set_run_font(r, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Versão 1.0")
    set_run_font(r, size=12, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(240)
    r = p.add_run("BRASIL\n2026")
    set_run_font(r, size=12)


def add_title_page(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("METALIQUE")
    set_run_font(r, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(85)
    r = p.add_run("MANUAL DO USUÁRIO\nSISTEMA METALIQUE INFINITY")
    set_run_font(r, size=18, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(7.5)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(85)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(
        "Manual operacional elaborado para orientar usuários e administradores no uso das telas, "
        "recursos e rotinas disponíveis no sistema Metalique Infinity."
    )
    set_run_font(r, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(210)
    r = p.add_run("BRASIL\n2026")
    set_run_font(r, size=12)


def add_toc(doc: Document, page_map: dict[str, int]) -> None:
    doc.add_page_break()
    h = doc.add_paragraph("SUMÁRIO", style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(18)

    items = [
        ("1 APRESENTAÇÃO", 0),
        ("2 ACESSO AO SISTEMA", 0),
        ("3 ÁREA AUTENTICADA", 0),
        ("4 MÓDULO QUALIDADE", 0),
        ("5 ADMINISTRAÇÃO DE USUÁRIOS", 0),
        ("6 SEGURANÇA E PRIMEIRO ACESSO", 0),
        ("7 SOLUÇÃO DE PROBLEMAS", 0),
    ]
    for label, level in items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(level * 0.75)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.line_spacing = 1.0
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Cm(15.5))
        r = p.add_run(label)
        set_run_font(r, size=12, bold=True)
        r = p.add_run("\t" + str(page_map.get(label, "--")))
        set_run_font(r, size=12)


def add_intro(doc: Document) -> None:
    h = doc.add_paragraph("1 APRESENTAÇÃO", style="Heading 1")
    h.paragraph_format.space_before = Pt(0)
    p = doc.add_paragraph(
        "Este manual descreve as telas existentes no Metalique Infinity na data de sua elaboração. "
        "O conteúdo foi organizado para apoiar a operação diária, o treinamento de novos usuários e "
        "a administração dos acessos. As opções visíveis variam conforme o perfil e as permissões da conta."
    )
    p.paragraph_format.first_line_indent = Cm(1.25)

    doc.add_paragraph("1.1 Público-alvo", style="Heading 2")
    for text in (
        "Colaboradores que consultam o Dashboard ou os indicadores da Qualidade.",
        "Responsáveis pelo registro de RAPs e produtos coletados.",
        "Administradores que criam contas e gerenciam permissões.",
    ):
        doc.add_paragraph(text, style="List Bullet")

    doc.add_paragraph("1.2 Convenções utilizadas", style="Heading 2")
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(11.8)
    headers = ("Convenção", "Significado")
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = (Cm(4.2), Cm(11.8))[i]
        set_cell_shading(cell, "EDEDED")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        set_run_font(r, size=11, bold=True)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Asterisco (*)", "Campo obrigatório para concluir um cadastro."),
        ("Botão vermelho", "Ação principal da tela, como entrar, gravar ou criar."),
        ("Botão contornado", "Ação secundária, como cancelar, filtrar ou imprimir."),
        ("Opção indisponível", "A ação depende de permissão, estado da conta ou preenchimento prévio."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        for i, value in enumerate((left, right)):
            cells[i].width = (Cm(4.2), Cm(11.8))[i]
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(value)
            set_run_font(r, size=10.5, bold=(i == 0))
    set_table_geometry(table, [2381, 6690], indent_dxa=140)


def add_label_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(label + ": ")
    set_run_font(r, size=11, bold=True, color=RED)
    r = p.add_run(text)
    set_run_font(r, size=11)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.right_indent = Cm(0.15)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_RED)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "DA0F0F")
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r = p.add_run("ATENÇÃO: ")
    set_run_font(r, size=10.5, bold=True, color=RED)
    r = p.add_run(text)
    set_run_font(r, size=10.5)


def set_image_alt_text(inline_shape, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_screen(doc: Document, screen: dict, figure_number: int) -> None:
    h = doc.add_paragraph(screen["title"], style="Heading 2")
    h.paragraph_format.page_break_before = True
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(7)

    add_label_paragraph(doc, "Objetivo", screen["purpose"])
    add_label_paragraph(doc, "Como acessar", screen["access"])

    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.first_line_indent = Cm(0)
    image_p.paragraph_format.space_before = Pt(5)
    image_p.paragraph_format.space_after = Pt(0)
    image_p.paragraph_format.keep_with_next = True
    shape = image_p.add_run().add_picture(str(ASSETS / screen["image"]), width=Cm(15.8))
    set_image_alt_text(shape, f"Captura da {screen['title']}")

    cap = doc.add_paragraph(style="Figure Caption")
    r = cap.add_run(f"Figura {figure_number} - {screen['title'].split(' ', 1)[1]}")
    set_run_font(r, size=10)
    src = doc.add_paragraph(style="Figure Source")
    r = src.add_run("Fonte: Elaborado pelo autor a partir do sistema Metalique Infinity (2026).")
    set_run_font(r, size=10, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("Procedimento")
    set_run_font(r, size=11, bold=True, color=RED)
    for step in screen["steps"]:
        doc.add_paragraph(step, style="List Number")
    add_note(doc, screen["note"])


def add_security_and_troubleshooting(doc: Document) -> None:
    h = doc.add_paragraph("6 SEGURANÇA E PRIMEIRO ACESSO", style="Heading 1")
    h.paragraph_format.page_break_before = True
    h.paragraph_format.space_before = Pt(0)
    for title, text in (
        ("6.1 Troca obrigatória de senha", "Contas criadas com senha inicial podem ser direcionadas para a troca obrigatória antes de acessar os módulos. Informe a senha temporária, cadastre uma nova senha válida, confirme-a e conclua a alteração."),
        ("6.2 Proteção da conta", "Não compartilhe credenciais, não mantenha senhas anotadas em locais visíveis e encerre a sessão pelo comando Sair. Em computadores compartilhados, confirme que a tela pública foi exibida antes de se afastar."),
        ("6.3 Permissões", "As permissões controlam tanto o que aparece na interface quanto o que o servidor aceita executar. Solicite ao administrador somente os acessos necessários para sua atividade."),
    ):
        doc.add_paragraph(title, style="Heading 2")
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.25)

    h = doc.add_paragraph("7 SOLUÇÃO DE PROBLEMAS", style="Heading 1")
    h.paragraph_format.page_break_before = True
    h.paragraph_format.space_before = Pt(0)
    problems = [
        ("Não consigo entrar", "Confira e-mail e senha, verifique se a conta está ativa e use a recuperação de senha quando necessário."),
        ("Uma tela ou botão não aparece", "A conta provavelmente não possui a permissão correspondente. Solicite a revisão ao administrador."),
        ("O módulo Qualidade não carrega", "Atualize a página, confira a conexão e tente novamente. Se persistir, registre a mensagem apresentada e acione o suporte."),
        ("Não consigo gravar um RAP", "Revise os campos obrigatórios, a descrição mínima e a seleção de colaboradores."),
        ("Não consigo gravar uma coleta", "Confirme os campos obrigatórios e a inclusão de uma a seis fotos válidas, com até 5 MB cada."),
        ("A sessão expirou", "Atualize a página, faça login novamente e repita a operação. Dados ainda não gravados podem precisar ser preenchidos outra vez."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    widths = (Cm(5.0), Cm(11.0))
    for i, value in enumerate(("Situação", "Orientação")):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, "EDEDED")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        set_run_font(r, size=11, bold=True)
    set_repeat_table_header(table.rows[0])
    for situation, orientation in problems:
        cells = table.add_row().cells
        for i, value in enumerate((situation, orientation)):
            cells[i].width = widths[i]
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(value)
            set_run_font(r, size=10.5, bold=(i == 0))
    set_table_geometry(table, [2835, 6236], indent_dxa=140)

    doc.add_paragraph("7.1 Informações para o suporte", style="Heading 2")
    p = doc.add_paragraph("Ao solicitar ajuda, informe:")
    p.paragraph_format.first_line_indent = Cm(0)
    for item in (
        "nome da tela e ação realizada;",
        "mensagem exibida;",
        "data e horário aproximados;",
        "identificador do RAP ou RETIR, quando aplicável;",
        "captura de tela sem expor senha ou informação sensível.",
    ):
        doc.add_paragraph(item, style="List Bullet")


def set_document_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Manual do Usuário - Metalique Infinity"
    props.subject = "Manual operacional das telas do sistema"
    props.author = "Metalique"
    props.keywords = "Metalique Infinity, manual do usuário, qualidade, sistema"
    props.comments = "Documento elaborado em formato ABNT para uso interno."


def build(page_map: dict[str, int]) -> None:
    doc = Document()
    configure_page(doc.sections[0], numbered=False)
    configure_styles(doc)
    set_document_properties(doc)

    add_cover(doc)
    add_title_page(doc)
    add_toc(doc, page_map)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(body_section, numbered=True)
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), "1")
    body_section._sectPr.append(pg_num_type)

    add_intro(doc)
    figure_number = 1
    for screen in SCREENS:
        if screen.get("section"):
            h = doc.add_paragraph(screen["section"], style="Heading 1")
            h.paragraph_format.page_break_before = True
            h.paragraph_format.space_before = Pt(0)
            p = doc.add_paragraph(
                {
                    "2 ACESSO AO SISTEMA": "Esta seção reúne as telas públicas de entrada, recuperação e solicitação de acesso.",
                    "3 ÁREA AUTENTICADA": "Esta seção apresenta a navegação comum às áreas internas e os recursos pessoais do cabeçalho.",
                    "4 MÓDULO QUALIDADE": "Esta seção descreve os indicadores, filtros, registros e formulários operacionais da Qualidade.",
                    "5 ADMINISTRAÇÃO DE USUÁRIOS": "Esta seção é destinada às contas com permissão para criar, editar e controlar acessos.",
                }[screen["section"]]
            )
            p.paragraph_format.first_line_indent = Cm(1.25)
        add_screen(doc, screen, figure_number)
        figure_number += 1

    add_security_and_troubleshooting(doc)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(OUTPUT)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--page-map", type=Path)
    args = parser.parse_args()
    page_map = {}
    if args.page_map and args.page_map.exists():
        page_map = json.loads(args.page_map.read_text(encoding="utf-8"))
    build(page_map)
    print(OUTPUT)


if __name__ == "__main__":
    main()
